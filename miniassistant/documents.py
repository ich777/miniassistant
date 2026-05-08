"""
Dokument-Extraktion fuer Anhaenge (PDF, DOCX, plain text).

Liefert Text fuer den LLM-Prompt + optional Seiten-PNGs (gescannte PDFs)
fuer den Vision-Pfad. Alle Deps sind optional (extra: docs).
"""
from __future__ import annotations

import base64
import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# MIME-Types die wir verarbeiten (zusaetzlich zu Bildern, die direkt an Vision gehen)
SUPPORTED_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
    "application/xml",
    "text/xml",
}

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xml", ".log", ".rst"}


def is_supported(mime: str, filename: str = "") -> bool:
    if mime and mime.lower() in SUPPORTED_MIMES:
        return True
    if mime and mime.lower().startswith("text/"):
        return True
    if filename:
        ext = Path(filename).suffix.lower()
        if ext in SUPPORTED_EXTS:
            return True
    return False


def extract_document(
    data: bytes,
    filename: str,
    mime: str = "",
    *,
    max_chars: int = 30000,
    max_pages_render: int = 10,
) -> dict[str, Any]:
    """
    Extrahiert Text und/oder Bilder aus einem Dokument-Anhang.

    Returns:
        {
          "name": str,
          "text": str,                 # Extrahierter Text (kann leer sein)
          "images": list[{mime,data}], # PNG-Seiten bei gescannten PDFs (Vision-Fallback)
          "pages": int,                # Gesamtseiten (PDF) oder 0
          "truncated": bool,
          "error": str | "",
        }
    """
    out: dict[str, Any] = {
        "name": filename or "anhang",
        "text": "",
        "images": [],
        "pages": 0,
        "truncated": False,
        "error": "",
    }
    mime_l = (mime or "").lower()
    ext = Path(filename or "").suffix.lower()
    try:
        if mime_l == "application/pdf" or ext == ".pdf":
            _extract_pdf(data, out, max_chars=max_chars, max_pages_render=max_pages_render)
        elif ext == ".docx" or "wordprocessingml" in mime_l:
            _extract_docx(data, out, max_chars=max_chars)
        else:
            # Plain text / markdown / csv / json / xml / log
            _extract_text(data, out, max_chars=max_chars)
    except Exception as e:
        logger.exception("Dokument-Extraktion fehlgeschlagen fuer %s", filename)
        out["error"] = str(e) or type(e).__name__
    return out


def _extract_text(data: bytes, out: dict[str, Any], *, max_chars: int) -> None:
    # Encoding-Detect: utf-8 first, fallback latin-1
    text = ""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if not text:
        text = data.decode("utf-8", errors="replace")
    if len(text) > max_chars:
        out["text"] = text[:max_chars]
        out["truncated"] = True
    else:
        out["text"] = text


def _extract_pdf(data: bytes, out: dict[str, Any], *, max_chars: int, max_pages_render: int) -> None:
    try:
        from pypdf import PdfReader
    except ImportError:
        out["error"] = "pypdf nicht installiert (pip install -e '.[docs]')"
        return
    import io
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        out["error"] = f"PDF nicht lesbar: {e}"
        return
    pages = len(reader.pages)
    out["pages"] = pages
    text_parts: list[str] = []
    total_chars = 0
    chars_per_page: list[int] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        chars_per_page.append(len(t.strip()))
        if total_chars + len(t) > max_chars:
            text_parts.append(t[: max(0, max_chars - total_chars)])
            out["truncated"] = True
            total_chars = max_chars
            break
        text_parts.append(t)
        total_chars += len(t)
    out["text"] = "\n\n".join(text_parts).strip()

    # Heuristik: <50 chars/Seite im Schnitt = wahrscheinlich Scan → Bilder rendern
    avg = (sum(chars_per_page) / max(1, len(chars_per_page))) if chars_per_page else 0
    if pages > 0 and avg < 50:
        rendered = _render_pdf_pages(data, max_pages=max_pages_render)
        if rendered:
            out["images"] = rendered
            if pages > max_pages_render:
                out["truncated"] = True
            # Hinweistext, damit LLM weiss warum Text leer ist
            note = f"[PDF gescannt — {min(pages, max_pages_render)}/{pages} Seiten als Bild beigelegt]"
            out["text"] = (out["text"] + "\n\n" + note).strip() if out["text"] else note


def _render_pdf_pages(data: bytes, *, max_pages: int) -> list[dict[str, Any]]:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        logger.warning("pypdfium2 nicht installiert — gescannte PDFs koennen nicht gerendert werden (pip install -e '.[docs]')")
        return []
    import io
    images: list[dict[str, Any]] = []
    try:
        pdf = pdfium.PdfDocument(data)
        n = min(len(pdf), max_pages)
        for i in range(n):
            page = pdf[i]
            # 144 DPI = scale 2.0 (PDF default 72 DPI). Genug fuer OCR via Vision-Model.
            pil_image = page.render(scale=2.0).to_pil()
            buf = io.BytesIO()
            pil_image.save(buf, format="PNG", optimize=True)
            b64 = base64.b64encode(buf.getvalue()).decode("ascii")
            images.append({"mime_type": "image/png", "data": b64})
    except Exception as e:
        logger.warning("PDF-Rendering fehlgeschlagen: %s", e)
    return images


def _extract_docx(data: bytes, out: dict[str, Any], *, max_chars: int) -> None:
    try:
        import docx  # python-docx
    except ImportError:
        out["error"] = "python-docx nicht installiert (pip install -e '.[docs]')"
        return
    import io
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:
        out["error"] = f"DOCX nicht lesbar: {e}"
        return
    parts: list[str] = []
    total = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        if total + len(t) > max_chars:
            parts.append(t[: max(0, max_chars - total)])
            out["truncated"] = True
            total = max_chars
            break
        parts.append(t)
        total += len(t) + 1
    # Tabellen mitnehmen (zeilenweise tab-separiert)
    if total < max_chars:
        for tbl in doc.tables:
            for row in tbl.rows:
                line = "\t".join((c.text or "").strip() for c in row.cells)
                if not line.strip():
                    continue
                if total + len(line) > max_chars:
                    out["truncated"] = True
                    total = max_chars
                    break
                parts.append(line)
                total += len(line) + 1
            if total >= max_chars:
                break
    out["text"] = "\n".join(parts).strip()


def format_document_block(doc: dict[str, Any]) -> str:
    """Wrappt extrahierten Text als <doc>-Block fuer den Prompt.
    Wird von chat_loop beim History-Save wieder gestrippt."""
    name = doc.get("name") or "anhang"
    text = doc.get("text") or ""
    if not text:
        return ""
    trunc = " (gekuerzt)" if doc.get("truncated") else ""
    return f"<doc name=\"{name}\"{trunc and ' truncated=\"true\"'}>\n{text}\n</doc>"


def fit_documents_to_budget(text: str, max_chars: int) -> str:
    """Kuerzt <doc>-Bloecke proportional, falls Gesamttext > max_chars.

    Trimmt nur Block-Inhalte, nie den User-Text drumherum. Markiert gekuerzte Bloecke
    mit truncated="true" + sichtbarem Hinweis. Wenn nicht-doc-Text allein schon > max_chars:
    text wird hart auf max_chars getrimmt (Notbremse).
    """
    if not text or len(text) <= max_chars:
        return text
    import re
    pattern = re.compile(r'(<doc\b[^>]*>)(.*?)(</doc>)', re.DOTALL)
    blocks = list(pattern.finditer(text))
    if not blocks:
        return text[:max_chars]  # keine docs zum Kuerzen
    non_doc_chars = len(text) - sum(len(m.group(0)) for m in blocks)
    overhead_per_block = 80  # opening tag + closing tag + Hinweistext
    available = max_chars - non_doc_chars - len(blocks) * overhead_per_block
    if available <= 0:
        return text[:max_chars]
    total_inner = sum(len(m.group(2)) for m in blocks)
    if total_inner <= available:
        return text
    ratio = available / total_inner

    def _shrink(m: "re.Match[str]") -> str:
        opening, inner, closing = m.group(1), m.group(2), m.group(3)
        new_len = max(0, int(len(inner) * ratio))
        if new_len < len(inner):
            inner = inner[:new_len].rstrip() + "\n[...gekuerzt um in Kontext zu passen...]"
            if 'truncated=' not in opening:
                opening = opening[:-1] + ' truncated="true">'
        return opening + inner + closing

    return pattern.sub(_shrink, text)


def strip_document_blocks(text: str) -> str:
    """Ersetzt <doc>...</doc> Bloecke durch kurze Marker (fuer History-Save).
    Spart Kontext-Tokens im Verlauf nachfolgender Turns."""
    if not text or "<doc " not in text:
        return text
    import re
    def _replace(m: "re.Match[str]") -> str:
        block = m.group(0)
        # Name und Text-Laenge extrahieren
        name_m = re.search(r'name="([^"]*)"', block)
        name = name_m.group(1) if name_m else "anhang"
        inner = re.sub(r'^<doc[^>]*>\s*|\s*</doc>$', '', block, flags=re.DOTALL)
        return f"[Anhang: {name} — {len(inner)} Zeichen]"
    return re.sub(r'<doc [^>]*>.*?</doc>', _replace, text, flags=re.DOTALL)
